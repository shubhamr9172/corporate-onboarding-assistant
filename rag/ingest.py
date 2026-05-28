import os
import sys
import logging
from typing import List

# Ensure project root is in the Python path (Issue #10)
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import chromadb

# Initialize logging
logger = logging.getLogger("app.ingest")


def load_text_or_markdown(file_path: str) -> List[Document]:
    """Loads plain text or markdown files."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        metadata = {
            "source": os.path.basename(file_path),
            "file_type": os.path.splitext(file_path)[1],
        }
        return [Document(page_content=text, metadata=metadata)]
    except Exception as e:
        logger.error(f"Error loading text file {file_path}: {e}")
        return []


def load_pdf(file_path: str) -> List[Document]:
    """Loads text from PDF files using pypdf."""
    try:
        import pypdf

        reader = pypdf.PdfReader(file_path)
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                metadata = {
                    "source": os.path.basename(file_path),
                    "file_type": ".pdf",
                    "page": i + 1,
                }
                docs.append(Document(page_content=text, metadata=metadata))
        return docs
    except ImportError:
        logger.error("pypdf library not found. Install requirements.")
        return []
    except Exception as e:
        logger.error(f"Error loading PDF file {file_path}: {e}")
        return []


def load_docx(file_path: str) -> List[Document]:
    """Loads text from Word (.docx) files using python-docx."""
    try:
        import docx

        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                full_text.append(para.text)

        # Parse tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    full_text.append(" | ".join(row_text))

        text = "\n".join(full_text)
        metadata = {"source": os.path.basename(file_path), "file_type": ".docx"}
        return [Document(page_content=text, metadata=metadata)]
    except ImportError:
        logger.error("python-docx library not found. Install requirements.")
        return []
    except Exception as e:
        logger.error(f"Error loading Word file {file_path}: {e}")
        return []


def determine_required_role(file_path: str) -> str:
    """Assigns role metadata required to access this file (Issue #2)."""
    filename = os.path.basename(file_path).lower()
    # Check parent folder names or filename prefix
    parts = file_path.lower().replace("\\", "/").split("/")
    if "admin" in parts or filename.startswith("admin_"):
        return "admin"
    if "hr" in parts or filename.startswith("hr_"):
        return "HR"
    return "joinee"


def scan_and_load_data(data_dir: str) -> List[Document]:
    """Scans data directory recursively and loads supported file types with role permissions."""
    all_documents = []
    if not os.path.exists(data_dir):
        logger.warning(f"Data directory {data_dir} does not exist.")
        return []

    for root, dirs, files in os.walk(data_dir):
        for filename in files:
            file_path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()
            logger.info(f"Processing file: {filename} ({ext}) in {root}")

            role = determine_required_role(file_path)
            file_docs = []

            if ext in [".txt", ".md"]:
                file_docs = load_text_or_markdown(file_path)
            elif ext == ".pdf":
                file_docs = load_pdf(file_path)
            elif ext == ".docx":
                file_docs = load_docx(file_path)
            else:
                logger.warning(f"Unsupported file format: {filename}")

            for d in file_docs:
                d.metadata["required_role"] = role

            all_documents.extend(file_docs)

    return all_documents


def run_ingestion():
    """
    Main ingestion execution block.
    Extracts text, splits it into chunks, embeddings, and stores in ChromaDB.
    """
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    dotenv_path = os.path.join(root_dir, ".env")
    load_dotenv(dotenv_path=dotenv_path)
    data_dir = os.path.join(root_dir, "data")
    persist_dir = os.path.join(root_dir, "chroma_db")

    logger.info("Starting document ingestion process...")

    # 2. Load Documents
    raw_docs = scan_and_load_data(data_dir)
    if not raw_docs:
        logger.warning("No documents found to ingest.")
        return

    logger.info(f"Loaded {len(raw_docs)} document files/pages.")

    # 3. Split Text
    # Target chunk sizes: 300 to 800 tokens.
    # ~1000 characters is roughly 250-300 words (350-400 tokens)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
    chunks = text_splitter.split_documents(raw_docs)
    logger.info(f"Split documents into {len(chunks)} text chunks.")

    # 4. Generate Embeddings & Store in ChromaDB
    google_key = os.getenv("GOOGLE_API_KEY")
    if not google_key:
        logger.error("GOOGLE_API_KEY is not configured. Ingestion aborted.")
        sys.exit(1)

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001", google_api_key=google_key
    )

    try:
        # Initialize chroma client using the singleton manager (Issue #10)
        from utils.chroma_manager import get_chroma_client

        chroma_client = get_chroma_client()

        # Get or create collection
        collection = chroma_client.get_or_create_collection(
            name="onboarding_faq_collection", metadata={"hnsw:space": "cosine"}
        )

        # Clear existing collection items to allow clean refreshes
        count = collection.count()
        if count > 0:
            logger.info(
                f"Pruning {count} existing records from collection onboarding_faq_collection."
            )
            # Clear by fetching all ids
            all_ids = collection.get()["ids"]
            if all_ids:
                collection.delete(ids=all_ids)

        # Bulk insert chunks
        ids = [f"chunk_{i}" for i in range(len(chunks))]
        documents = [c.page_content for c in chunks]
        metadatas = [c.metadata for c in chunks]

        logger.info("Generating embeddings and writing chunks to database...")
        # langchain embeddings client used to pre-vectorize or let Chroma handle it
        vectors = embeddings.embed_documents(documents)

        # Add to collection
        collection.add(
            ids=ids, embeddings=vectors, documents=documents, metadatas=metadatas
        )

        logger.info(
            f"Ingestion completed. Stored {collection.count()} chunks in ChromaDB."
        )

    except Exception as e:
        logger.error(f"Error during ChromaDB write operations: {e}")
        sys.exit(1)


if __name__ == "__main__":
    # Setup simple console logging if run independently
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] (%(module)s) - %(message)s",
    )
    run_ingestion()
