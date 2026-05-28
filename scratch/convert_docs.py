import os
import sys
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


def create_server_db_docx(output_path):
    doc = docx.Document()
    doc.add_heading("Corporate Onboarding: Server & Database Access Guide", level=1)

    doc.add_heading("1. Security Compliance & Prerequisites", level=2)
    doc.add_paragraph(
        "Before requesting access to any server or database, you must complete the following mandatory security training courses:"
    )
    doc.add_paragraph(
        "- Securing Remote Access & SSH Hygiene (Course Code: SEC-SSH-101)"
    )
    doc.add_paragraph("- Data Protection & Database Security (Course Code: SEC-DB-202)")
    doc.add_paragraph(
        "Additionally, access to production environments is strictly audited and subject to Least Privilege principles."
    )

    doc.add_heading("2. Server Access (Linux & Windows)", level=2)
    doc.add_paragraph(
        "Access to corporate servers is restricted and must be routed through secure bastion hosts."
    )
    doc.add_heading("A. Non-Production Servers (Dev / Staging)", level=3)
    doc.add_paragraph(
        "Access Protocol: SSH (for Linux) or RDP (for Windows) via Bastion Host.\n"
        "Request Platform: ServiceNow Portal -> Request Catalog -> 'Server Access - Non-Prod'\n"
        "Required Information:\n"
        "- Target server hostnames or IP addresses.\n"
        "- Public SSH Key (id_rsa.pub or id_ed25519.pub).\n"
        "- Justification for access.\n"
        "Approval Flow: Line Manager approval."
    )

    doc.add_heading("B. Production Servers", level=3)
    doc.add_paragraph(
        "Access Protocol: Session-based SSH/RDP brokered by CyberArk Privileged Access Manager (PAM).\n"
        "Request Platform: SailPoint Identity Portal -> 'Production Server access'\n"
        "Approval Flow: Line Manager approval and System Owner approval.\n"
        "Note: Direct root or Administrator logins are blocked. All actions are logged and audited."
    )

    doc.add_heading("3. Database Access (SQL & NoSQL)", level=2)
    doc.add_paragraph(
        "Direct database connection requests must be justified and specify the database schema and read/write permission levels."
    )
    doc.add_heading("A. Development Databases", level=3)
    doc.add_paragraph(
        "Access Protocol: Local connections or direct VPN connection using individual credentials.\n"
        "Request Platform: ServiceNow -> 'DB Access Request'\n"
        "Supported DBs: PostgreSQL, MySQL, MongoDB, Oracle.\n"
        "Approval Flow: Technical Lead approval."
    )

    doc.add_heading("B. Production Databases", level=3)
    doc.add_paragraph(
        "Access Protocol: Temporary connection strings or database sessions managed via IAM Database Authentication (AWS IAM / Azure Active Directory) or CyberArk.\n"
        "Strict Rule: Production database write access is strictly limited to CI/CD migration scripts. Direct write queries by developers are prohibited.\n"
        "Approval Flow: Line Manager, Database Administrator (DBA), and Chief Information Security Officer (CISO) approvals are required."
    )

    doc.save(output_path)
    print(f"Created docx at: {output_path}")


def create_privileged_accounts_pdf(output_path):
    pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(
        Paragraph(
            "Corporate Onboarding: Privileged & Administrator Accounts", styles["Title"]
        )
    )
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Definition of Privileged Access", styles["Heading2"]))
    story.append(Spacer(1, 6))
    story.append(
        Paragraph(
            "Privileged access refers to administrative rights that allow users to change system configurations, bypass security controls, or access sensitive system directories. Examples include:",
            styles["Normal"],
        )
    )
    story.append(
        Paragraph(
            "- Local Administrator: Administrative rights on your issued corporate laptop.<br/>"
            "- Domain Administrator: Management access across the Active Directory domain.<br/>"
            "- Root/Superuser: Access to systems or servers at the highest privilege level.<br/>"
            "- Service Accounts: Non-human accounts used by software applications.",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 12))

    story.append(
        Paragraph("2. Temporary Local Administrator Access", styles["Heading2"])
    )
    story.append(Spacer(1, 6))
    story.append(
        Paragraph(
            "To install specialized software or modify system network interfaces on your company-issued device, use the MakeMeAdmin app (Windows) or Privileges app (macOS). Privilege escalation is granted for 30 minutes per session.",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            "3. Persistent Privileged Accounts (Domain / Server Admin)",
            styles["Heading2"],
        )
    )
    story.append(Spacer(1, 6))
    story.append(
        Paragraph(
            "Secondary Admin Accounts: You must request a separate account prefixed with adm- (e.g., adm-jdoe). Request via SailPoint Identity Portal -> 'Create Administrative Account'.",
            styles["Normal"],
        )
    )
    story.append(
        Paragraph(
            "CyberArk PAM: For root and database admin access, credentials are managed through CyberArk PVWA console at https://cyberark-vault.corp.internal.",
            styles["Normal"],
        )
    )

    pdf_doc.build(story)
    print(f"Created pdf at: {output_path}")


if __name__ == "__main__":
    data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data"))
    os.makedirs(data_dir, exist_ok=True)

    # 1. Create Server and DB access as a DOCX file
    create_server_db_docx(os.path.join(data_dir, "access_servers_databases.docx"))

    # 2. Create Privileged Accounts as a PDF file
    create_privileged_accounts_pdf(
        os.path.join(data_dir, "access_privileged_accounts.pdf")
    )

    # 3. Clean up the old MD files if they were replaced by these formats to avoid duplication
    old_md_files = [
        os.path.join(data_dir, "access_servers_databases.md"),
        os.path.join(data_dir, "access_privileged_accounts.md"),
    ]
    for old_file in old_md_files:
        if os.path.exists(old_file):
            os.remove(old_file)
            print(f"Removed old duplicate: {old_file}")
